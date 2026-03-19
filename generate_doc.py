#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_deployment_guide():
    """创建个人网站构建和部署指南 Word 文档"""
    
    doc = Document()
    
    # 设置中文字体支持
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题样式
    title = doc.add_heading('个人网站构建与部署指南', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('从零开始搭建并部署你的个人介绍网站')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # 空行
    
    # 目录信息
    info = doc.add_paragraph()
    info.add_run('作者：AN4FJWS\n').bold = True
    info.add_run('GitHub: https://github.com/an4fjws/personal_website\n')
    info.add_run('在线访问：https://an4fjws.github.io/personal_website/\n')
    info.add_run('生成日期：2026 年 3 月')
    
    doc.add_page_break()
    
    # ========== 第一章 ==========
    doc.add_heading('第一章：项目概述', level=1)
    
    doc.add_heading('1.1 项目简介', level=2)
    doc.add_paragraph(
        '本项目是一个现代化的响应式个人介绍网站，包含首页、关于我、技能展示、'
        '项目展示和联系方式等模块。网站采用纯 HTML/CSS/JavaScript 构建，无需任何框架依赖，'
        '可以免费部署在 GitHub Pages 上。'
    )
    
    doc.add_heading('1.2 技术栈', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('HTML5').bold = True
    p.add_run(' - 语义化页面结构')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('CSS3').bold = True
    p.add_run(' - 响应式设计、渐变背景、动画效果')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('JavaScript').bold = True
    p.add_run(' - 原生 JS，无第三方依赖')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('GitHub Pages').bold = True
    p.add_run(' - 免费静态网站托管')
    
    doc.add_heading('1.3 功能特性', level=2)
    features = [
        ('响应式设计', '适配桌面、平板、手机等各种设备'),
        ('平滑滚动', '导航栏点击平滑滚动到对应区域'),
        ('移动端菜单', '汉堡菜单适配小屏幕'),
        ('滚动动画', '元素进入视口时触发动画'),
        ('回到顶部', '滚动后显示回到顶部按钮'),
        ('联系表单', '内置联系表单（需后端支持）'),
    ]
    
    for feature, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(feature).bold = True
        p.add_run(f' - {desc}')
    
    doc.add_page_break()
    
    # ========== 第二章 ==========
    doc.add_heading('第二章：本地开发环境', level=1)
    
    doc.add_heading('2.1 前置要求', level=2)
    doc.add_paragraph('请确保你的计算机已安装以下工具：')
    
    requirements = [
        ('Git', '版本控制工具，https://git-scm.com/'),
        ('文本编辑器', 'VS Code、Sublime Text 或任意编辑器'),
        ('现代浏览器', 'Chrome、Firefox、Safari 等'),
    ]
    
    for tool, desc in requirements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(tool).bold = True
        p.add_run(f' - {desc}')
    
    doc.add_heading('2.2 获取项目代码', level=2)
    doc.add_paragraph('方法一：从 GitHub 克隆（如果已有仓库）')
    
    doc.add_paragraph('git clone https://github.com/an4fjws/personal_website.git', style='No Spacing')
    doc.add_paragraph('cd personal_website', style='No Spacing')
    
    doc.add_paragraph('方法二：直接下载 ZIP 文件')
    doc.add_paragraph('访问 https://github.com/an4fjws/personal_website，点击 "Code" → "Download ZIP"，解压即可。')
    
    doc.add_heading('2.3 本地预览', level=2)
    
    doc.add_paragraph('方法一：直接打开')
    doc.add_paragraph('双击 index.html 文件，用浏览器打开即可预览。')
    
    doc.add_paragraph('方法二：使用本地服务器（推荐）')
    
    p = doc.add_paragraph()
    p.add_run('Python 3:').bold = True
    code = doc.add_paragraph('python3 -m http.server 8080', style='No Spacing')
    
    p = doc.add_paragraph()
    p.add_run('Node.js:').bold = True
    doc.add_paragraph('npx serve', style='No Spacing')
    
    doc.add_paragraph('然后在浏览器访问：http://localhost:8080')
    
    doc.add_page_break()
    
    # ========== 第三章 ==========
    doc.add_heading('第三章：自定义内容', level=1)
    
    doc.add_heading('3.1 修改个人信息', level=2)
    doc.add_paragraph('打开 index.html 文件，搜索并替换以下内容：')
    
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    header = info_table.rows[0].cells
    header[0].text = '搜索内容'
    header[1].text = '替换为'
    header[0].paragraphs[0].runs[0].bold = True
    header[1].paragraphs[0].runs[0].bold = True
    
    replacements = [
        ('张三', '你的姓名'),
        ('zhangsan@email.com', '你的邮箱'),
        ('+86 138 0000 0000', '你的电话'),
        ('北京市朝阳区', '你的所在地'),
        ('全栈开发工程师', '你的职业/头衔'),
    ]
    
    for search, replace in replacements:
        row = info_table.add_row()
        row.cells[0].text = search
        row.cells[1].text = replace
    
    doc.add_heading('3.2 修改技能列表', level=2)
    doc.add_paragraph('在 index.html 中找到 <section id="skills">，修改技能卡片内容：')
    doc.add_paragraph('每个技能卡片包含：技能图标、技能分类标题、技能列表项。')
    
    doc.add_heading('3.3 修改项目展示', level=2)
    doc.add_paragraph('在 index.html 中找到 <section id="projects">，修改：')
    doc.add_paragraph('项目名称、项目描述、技术标签、项目链接（GitHub/在线演示）。')
    
    doc.add_heading('3.4 修改社交链接', level=2)
    doc.add_paragraph('在 index.html 中找到 <section id="contact">，修改社交链接：')
    
    social = [
        ('GitHub', 'https://github.com/你的用户名'),
        ('LinkedIn', 'https://linkedin.com/in/你的用户名'),
        ('知乎', 'https://zhihu.com/people/你的用户名'),
    ]
    
    for platform, url in social:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{platform}: {url}')
    
    doc.add_page_break()
    
    # ========== 第四章 ==========
    doc.add_heading('第四章：部署到 GitHub Pages', level=1)
    
    doc.add_heading('4.1 创建 GitHub 账号', level=2)
    doc.add_paragraph('如果没有 GitHub 账号，请访问 https://github.com 注册。')
    
    doc.add_heading('4.2 创建 GitHub 仓库', level=2)
    
    steps = [
        '登录 GitHub：https://github.com',
        '点击右上角 + 按钮，选择 New repository，或直接访问 https://github.com/new',
        '填写仓库信息：',
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    
    sub_items = [
        'Repository name: personal_website',
        'Description（可选）: 个人介绍网站',
        'Public: 选择公开（免费部署需要公开仓库）',
        'Initialize this repository with a README: 不要勾选',
    ]
    
    for item in sub_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('点击 Create repository 按钮创建仓库。', style='List Number')
    
    doc.add_heading('4.3 配置 Git 环境', level=2)
    
    doc.add_paragraph('打开终端（Terminal），执行以下命令：')
    
    doc.add_paragraph('1. 配置 Git 用户名和邮箱：', style='List Number')
    doc.add_paragraph('git config --global user.name "你的用户名"\ngit config --global user.email "你的邮箱@example.com"', style='No Spacing')
    
    doc.add_paragraph('2. 初始化本地 Git 仓库：', style='List Number')
    doc.add_paragraph('cd /path/to/personal_website\ngit init', style='No Spacing')
    
    doc.add_paragraph('3. 添加并提交文件：', style='List Number')
    doc.add_paragraph('git add .\ngit commit -m "Initial commit"', style='No Spacing')
    
    doc.add_paragraph('4. 重命名分支为 main：', style='List Number')
    doc.add_paragraph('git branch -M main', style='No Spacing')
    
    doc.add_page_break()
    
    doc.add_heading('4.4 创建 Personal Access Token', level=2)
    doc.add_paragraph('GitHub 已不再支持使用密码进行 Git 操作，需要使用 Personal Access Token：')
    
    token_steps = [
        '访问 https://github.com/settings/tokens',
        '点击 Generate new token → Generate new token (classic)',
        '填写信息：',
    ]
    
    for step in token_steps:
        doc.add_paragraph(step, style='List Number')
    
    token_info = [
        'Note: Git Push（或其他易识别的名称）',
        'Expiration: No expiration 或 90 days',
        'Select scopes: 勾选 repo（全选）',
    ]
    
    for info in token_info:
        doc.add_paragraph(info, style='List Bullet')
    
    doc.add_paragraph('点击 Generate token，复制生成的 token（以 ghp_ 开头，只显示一次，请妥善保存！）', style='List Number')
    
    doc.add_heading('4.5 推送代码到 GitHub', level=2)
    
    doc.add_paragraph('方法一：使用 Personal Access Token（推荐）')
    doc.add_paragraph('在终端执行：', style='List Number')
    
    code_block = '''cd /path/to/personal_website
git remote add origin https://你的用户名:YOUR_TOKEN@github.com/你的用户名/personal_website.git
git push -u origin main'''
    doc.add_paragraph(code_block, style='No Spacing')
    
    doc.add_paragraph('将 YOUR_TOKEN 替换为你刚才复制的 token，将 你的用户名 替换为你的实际 GitHub 用户名。')
    
    doc.add_paragraph('推送时如果需要输入：', style='List Number')
    doc.add_paragraph('Username: 你的 GitHub 用户名\nPassword: 粘贴你的 Personal Access Token', style='No Spacing')
    
    doc.add_paragraph('方法二：使用 SSH（如果已配置 SSH 密钥）')
    doc.add_paragraph('git remote add origin git@github.com:你的用户名/personal_website.git\ngit push -u origin main', style='No Spacing')
    
    doc.add_page_break()
    
    doc.add_heading('4.6 启用 GitHub Pages', level=2)
    
    pages_steps = [
        '进入你的仓库页面：https://github.com/你的用户名/personal_website',
        '点击顶部的 Settings 标签',
        '在左侧菜单栏找到并点击 Pages',
        '在 Build and deployment 区域配置：',
    ]
    
    for step in pages_steps:
        doc.add_paragraph(step, style='List Number')
    
    pages_info = [
        'Source: 选择 Deploy from a branch',
        'Branch: 选择 main，文件夹选择 / (root)',
    ]
    
    for info in pages_info:
        doc.add_paragraph(info, style='List Bullet')
    
    doc.add_paragraph('点击 Save 按钮保存设置。', style='List Number')
    
    doc.add_heading('4.7 访问你的网站', level=2)
    doc.add_paragraph('等待 1-2 分钟部署完成后，访问：')
    
    url_para = doc.add_paragraph()
    url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url_run = url_para.add_run('https://你的用户名.github.io/personal_website/')
    url_run.font.size = Pt(14)
    url_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph('部署状态可以在仓库的 Actions 标签页查看。')
    
    doc.add_page_break()
    
    # ========== 第五章 ==========
    doc.add_heading('第五章：后续维护', level=1)
    
    doc.add_heading('5.1 更新网站内容', level=2)
    doc.add_paragraph('修改文件后，推送更新：')
    
    doc.add_paragraph('git add .\ngit commit -m "更新说明"\ngit push', style='No Spacing')
    
    doc.add_paragraph('GitHub Pages 会自动重新部署，约 1 分钟后生效。')
    
    doc.add_heading('5.2 绑定自定义域名（可选）', level=2)
    doc.add_paragraph('如果你想使用自己的域名（如 www.yourname.com）：')
    
    domain_steps = [
        '在域名注册商处添加 CNAME 记录，指向 你的用户名.github.io',
        '在仓库根目录创建 CNAME 文件，内容为你的域名（如：www.yourname.com）',
        '在仓库 Settings → Pages → Custom domain 中填写你的域名',
        '点击 Save 保存',
    ]
    
    for step in domain_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('5.3 查看部署日志', level=2)
    doc.add_paragraph('访问 https://github.com/你的用户名/personal_website/actions 查看部署历史和日志。')
    
    doc.add_heading('5.4 安全提示', level=2)
    doc.add_paragraph('⚠️ 重要：Personal Access Token 只在首次推送时使用，建议：')
    
    security_tips = [
        '不要将 Token 提交到代码仓库',
        'Token 泄露后立即删除并重新生成',
        '可以使用 Git 凭据管理器保存 Token',
        '考虑使用 SSH 密钥代替 Token',
    ]
    
    for tip in security_tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 第六章 ==========
    doc.add_heading('第六章：常见问题', level=1)
    
    faq = [
        ('Q: 推送时提示 "Repository not found"', 
         'A: 确保已在 GitHub 创建仓库，且仓库名称正确。检查远程仓库地址：git remote -v'),
        
        ('Q: 推送时认证失败', 
         'A: 确保使用 Personal Access Token 而非密码。Token 需要 repo 权限。'),
        
        ('Q: GitHub Pages 显示 404', 
         'A: 等待 1-2 分钟部署完成。检查 Settings → Pages 配置是否正确。确保仓库是公开的。'),
        
        ('Q: 修改后网站没有更新', 
         'A: 清除浏览器缓存（Ctrl+Shift+Delete 或 Cmd+Shift+Delete）。检查是否已成功推送：git log'),
        
        ('Q: 如何删除已部署的网站', 
         'A: 在 Settings → Pages 中将 Source 改为 None，或删除 GitHub 仓库。'),
    ]
    
    for question, answer in faq:
        p = doc.add_paragraph()
        p.add_run(question).bold = True
        doc.add_paragraph(answer)
    
    doc.add_page_break()
    
    # ========== 附录 ==========
    doc.add_heading('附录', level=1)
    
    doc.add_heading('A. 项目文件结构', level=2)
    doc.add_paragraph('''personal_website/
├── index.html      # 主页面（HTML 结构）
├── styles.css      # 样式文件（CSS 样式）
├── script.js       # 交互脚本（JavaScript 功能）
├── README.md       # 项目说明文档
├── .gitignore      # Git 忽略文件
└── generate_doc.py # 文档生成脚本''', style='No Spacing')
    
    doc.add_heading('B. 有用链接', level=2)
    links = [
        ('项目仓库', 'https://github.com/an4fjws/personal_website'),
        ('在线演示', 'https://an4fjws.github.io/personal_website/'),
        ('GitHub Pages 文档', 'https://pages.github.com/'),
        ('Git 下载', 'https://git-scm.com/'),
        ('VS Code', 'https://code.visualstudio.com/'),
    ]
    
    for name, url in links:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}: ').bold = True
        p.add_run(url)
    
    doc.add_heading('C. 技术支持', level=2)
    doc.add_paragraph('如有问题，请：')
    doc.add_paragraph('1. 查看本指南是否已解答你的问题', style='List Number')
    doc.add_paragraph('2. 在 GitHub 仓库提 Issue', style='List Number')
    doc.add_paragraph('3. 查阅官方文档：https://docs.github.com/en/pages', style='List Number')
    
    # 添加页脚
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('---\n').italic = True
    footer.add_run('文档生成完成。祝你部署顺利！\n')
    footer.add_run('如有问题，请访问：https://github.com/an4fjws/personal_website')
    
    # 保存文档
    doc.save('个人网站构建与部署指南.docx')
    print('✅ Word 文档已生成：个人网站构建与部署指南.docx')

if __name__ == '__main__':
    create_deployment_guide()
